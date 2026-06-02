from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面边距 ──
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# ── 字体工具 ──
FONT_BODY   = "微软雅黑"
FONT_CODE   = "Courier New"
COLOR_TITLE = RGBColor(0x1A, 0x56, 0xDB)   # 蓝
COLOR_H2    = RGBColor(0x1E, 0x3A, 0x5F)   # 深蓝
COLOR_H3    = RGBColor(0x1E, 0x3A, 0x5F)
COLOR_QA    = RGBColor(0x0F, 0x52, 0x2E)   # 深绿
COLOR_CODE  = RGBColor(0x1F, 0x2D, 0x3D)
COLOR_BG    = RGBColor(0xF3, 0xF4, 0xF6)   # 代码块背景（通过段落底纹实现）

def set_font(run, name=FONT_BODY, size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=22, bold=True, color=COLOR_TITLE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(16)

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=COLOR_H2)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # 下划线装饰
    run.underline = False
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A56DB")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=COLOR_H3)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def add_body(text, indent=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    """带 • 的项目符号段落（不依赖 Word 列表样式）"""
    p = doc.add_paragraph()
    prefix = "•  " if level == 0 else "  ◦  "
    run = p.add_run(prefix + text)
    set_font(run, size=11)
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_after   = Pt(4)

def add_code_block(lines):
    """灰色底纹 + 等宽字体代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    # 段落底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run("\n")
        run = p.add_run(line)
        set_font(run, name=FONT_CODE, size=10, color=COLOR_CODE)
    return p

def add_qa(q, a_lines):
    """问答块：Q 绿色加粗，A 普通"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    run_q = p.add_run("Q: " + q)
    set_font(run_q, size=11, bold=True, color=COLOR_QA)
    doc.add_paragraph()  # 占位
    # 删掉占位，直接加 A
    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(10)
    run_a = p2.add_run("A: ")
    set_font(run_a, size=11, bold=True)
    for i, line in enumerate(a_lines):
        if i > 0:
            p2.add_run("\n")
        run = p2.add_run(line)
        set_font(run, size=11)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # 蓝色背景
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A56DB")
        tcPr.append(shd)
    # 数据行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10.5)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)

def add_sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════════
#  正文内容
# ════════════════════════════════════════════

add_title("MySpringAIAgent 项目亮点与面试要点")

# 副标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Spring Boot 3.5 + Vue 3 + DeepSeek / Qwen | 全栈 AI 对话应用")
set_font(run, size=11, color=RGBColor(0x60, 0x74, 0x8A), italic=True)
p.paragraph_format.space_after = Pt(20)

# ─── 一、项目概述 ───
add_h2("一、项目概述")
add_body(
    "本项目是一个具备生产级功能的全栈 AI 对话平台，前端采用 Vue 3 + Vite，"
    "后端采用 Spring Boot 3.5 + Spring AI，通过 DeepSeek / 阿里云 Qwen 大模型提供智能对话能力，"
    "并集成 RAG 知识库、Redis 缓存、函数调用、SSE 流式输出等多项核心技术。"
)

# ─── 二、核心技术亮点 ───
add_h2("二、核心技术亮点")

# 2.1 SSE 流式对话
add_h3("2.1  全链路 SSE 流式对话（含中断恢复）")
add_bullet("前端使用 fetch + ReadableStream 解析 text/event-stream，实现打字机效果")
add_bullet("后端通过 AsyncContext 保持长连接，主动推送 LLM token 片段")
add_bullet("支持中断恢复：用户断线后可调用 continueStream() 从未完成处续写")
add_bullet("支持重新生成：regenerateStream() 删除旧回复、重新调用 LLM，temperature / maxTokens 实时生效")
add_bullet("支持 Reasoning Chunk：DeepSeek 的 <think> 思考内容单独渲染，可折叠展示")

add_qa(
    "SSE 和 WebSocket 的区别，为什么选 SSE？",
    [
        "SSE 是单向（服务端→客户端），天然适合 LLM token 流输出场景；基于标准 HTTP，",
        "无需协议升级握手，对反向代理、CDN、防火墙更友好。",
        "WebSocket 是全双工，适合双向高频通信（如在线游戏、协作编辑）。",
        "本项目对话写入由 REST POST 完成，只需单向推送，SSE 是更轻量的正确选择。",
    ]
)

# 2.2 RAG
add_h3("2.2  RAG 知识库（混合检索 + 时效衰减）")
add_body("完整数据链路：")
add_code_block([
    "文件上传 → Apache Tika 解析 → 按 800 token 分块",
    "→ DashScope Embedding（1536 维）",
    "→ Elasticsearch 存储（text 字段 + dense_vector 字段）",
    "──── 查询时 ────",
    "→ BM25 全文检索 + kNN 向量检索（混合搜索）",
    "→ 超取 15 个候选（TOP_K × 3）",
    "→ recencyMultiplier() 时效衰减重排",
    "   （7 天内 ×1.0，线性衰减至 90 天时 ×0.5）",
    "→ 取 Top-5 注入系统提示词",
])
add_bullet("MD5 增量更新：Scheduler 定时扫描文件目录，仅对变更文件重新 Embedding，未变更跳过")
add_bullet("混合搜索：BM25 精确词命中 + kNN 语义理解互补，召回率优于单一策略")
add_bullet("时效衰减：recencyMultiplier() 对老文档降权，避免旧知识优先级高于新内容")
add_bullet("并发安全：loadFromDirectory() 加 synchronized，防止 Scheduler 与 HTTP 触发重复写入")

add_qa(
    "BM25 和向量检索各有什么优缺点，为什么要混合？",
    [
        "BM25（TF-IDF 改进版）：对精确词（专有名词、产品型号、错别字）命中率高，",
        "不理解同义词和近义词；向量检索：理解语义，处理同义/近义表达，",
        "但对精确词不敏感。混合检索取二者之长，是 RAG 生产落地的标准方案，",
        "Elasticsearch 通过 RRF（Reciprocal Rank Fusion）算法融合两路结果。",
    ]
)

add_qa(
    "RAG 的效果如何评估？",
    [
        "项目层面：功能性验证（知识确认被检索到并注入回答）。",
        "生产级指标：Recall@K（Top-K 中是否包含相关文档）；",
        "RAGAS 框架三指标：faithfulness（答案是否忠于检索内容）、",
        "answer_relevancy（答案是否切题）、context_precision（检索内容是否精准）；",
        "也可人工标注问答对后做自动化回归测试。",
    ]
)

# 2.3 Spring AOP
add_h3("2.3  Spring AOP 代理陷阱的正确处理")
add_body(
    "KnowledgeService.loadFromDirectory() 在同类内部调用 this.deleteDocument() 和 "
    "this.uploadDocumentBytes()，导致 @Transactional 和 @CacheEvict 注解全部失效——"
    "调用的是原始对象而非 Spring 代理对象。"
)
add_body("解决方案：使用 @Lazy 自注入获得代理引用：", indent=0.5)
add_code_block([
    "@Autowired @Lazy",
    "private KnowledgeService self;   // 拿到 Spring 代理",
    "",
    "// 内部调用改为：",
    "self.deleteDocument(id);",
    "self.uploadDocumentBytes(bytes, name);",
])

add_qa(
    "Spring @Transactional 在哪些情况下会失效？",
    [
        "① 同类内部 this.method() 调用绕过 AOP 代理（本项目遇到的问题）；",
        "② 方法非 public（CGLIB 代理要求 public）；",
        "③ 异常被 catch 吞掉，Spring 无法感知需要回滚；",
        "④ Bean 不由 Spring 管理（new 出来的实例）；",
        "⑤ 数据库存储引擎不支持事务（如 MySQL MyISAM）。",
    ]
)

# 2.4 Redis 缓存
add_h3("2.4  Redis 分层缓存设计")
add_body("按数据特征设定差异化 TTL，避免同时过期引发缓存雪崩：")
add_table(
    ["缓存名", "TTL", "缓存内容", "说明"],
    [
        ["prompt_templates",   "24 小时", "提示词模板列表",   "写少读多，长 TTL 合理"],
        ["workflow_templates", "24 小时", "工作流模板",       "同上"],
        ["daily_questions",   "48 小时", "每日 AI 推荐问题", "每日生成一次，缓存跨天"],
        ["shared_conversations", "10 分钟", "分享链接内容",  "撤销时精准 evict(token)"],
        ["knowledge_documents", "24 小时", "知识库文档列表", "变更时手动 evict"],
    ]
)
add_bullet("CacheErrorHandler：Redis 故障时自动降级走 DB，主链路不受影响")
add_bullet("unless 条件：unless = \"#result == null || #result.isEmpty()\" 防止空结果缓存（缓存穿透）")
add_bullet("Jackson 序列化陷阱：List.of() 返回 ImmutableCollections，默认 Jackson 反序列化报错，通过 unless 条件杜绝此类对象入缓存")

add_qa(
    "缓存穿透、缓存雪崩、缓存击穿分别是什么，项目如何应对？",
    [
        "穿透：查询不存在的 key，每次都打穿到 DB。",
        "  → 项目用 unless 不缓存空/null 结果；生产可加布隆过滤器。",
        "雪崩：大量 key 同时过期，流量同时涌入 DB。",
        "  → 各缓存 TTL 错开（10min / 24h / 48h）；生产可加随机 jitter。",
        "击穿：单个热点 key 过期瞬间大量并发。",
        "  → 本项目场景并发低影响小；生产可用 Redisson 分布式锁或逻辑过期。",
    ]
)

# 2.5 Function Calling
add_h3("2.5  Function Calling / Tool Use（联网搜索 + 本地工具）")
add_bullet("Spring AI @Tool 注解注册工具函数，LLM 自主决定是否调用及传参")
add_bullet("ToolFunctions：查天气、查时间、计算器等本地工具")
add_bullet("WebSearchTools：集成 Bocha AI 搜索 API，实时联网检索最新信息")
add_bullet("MCP 工具：通过 spring-ai-starter-mcp-client 接入文件系统 + Brave Search，无需改业务代码")

add_qa(
    "Function Calling 的完整工作流程是什么？",
    [
        "① 第一轮请求：把工具的 JSON Schema（函数名 + 参数描述 + 返回描述）随对话一起发给 LLM；",
        "② LLM 判断：如需调用工具，返回结构化 tool_call（含函数名 + 参数 JSON），不返回自然语言；",
        "③ 应用层执行：解析 tool_call，调用对应函数，获取真实结果；",
        "④ 第二轮请求：将工具结果拼入对话历史，再次发给 LLM；",
        "⑤ LLM 基于工具结果生成最终自然语言回答。",
        "整个过程对用户完全透明，LLM 不直接执行代码，只负责「决策调用哪个工具」。",
    ]
)

# 2.6 多模型
add_h3("2.6  多模型动态切换（ChatClientRegistry）")
add_bullet("ChatClientRegistry 统一管理 DeepSeek、Qwen 等多个 ChatClient Bean")
add_bullet("前端通过 model 参数动态选择，后端按 modelId 路由到对应 Client")
add_bullet("业务代码只依赖 ChatClient 接口，新增模型只需配置 application.yml，不改业务逻辑")

# ─── 三、Redis 序列化深坑 ───
add_h2("三、工程细节：踩坑与解决")

add_h3("坑 1：ImmutableCollections 导致 Redis 反序列化崩溃")
add_body(
    "DailyQuestionController 使用 @Cacheable 缓存 AI 返回的问题列表。"
    "当 AI 调用失败时返回 List.of()（ImmutableCollections），"
    "Jackson DefaultTyping 将其序列化为裸数组 []（无类型信息）；"
    "下次读取时反序列化期望 [\"typename\",[...]] 格式，抛 MismatchedInputException。"
)
add_body("解决：unless = \"#result == null || #result.isEmpty()\"，空列表不入缓存，两个 bug 同时修复。")

add_h3("坑 2：Lettuce 连接池静默失效")
add_body(
    "application.yml 配置了 spring.data.redis.lettuce.pool.*，但实际连接池未启用。"
    "原因：Lettuce 连接池依赖 commons-pool2，未引入时配置被静默忽略。"
    "修复：pom.xml 加入 commons-pool2 依赖。"
)

add_h3("坑 3：@Scheduled + HTTP 并发写入重复行")
add_body(
    "KnowledgeRefreshScheduler 的定时任务与 /api/knowledge/refresh 接口并发触发时，"
    "loadFromDirectory() 可能被同时执行，对同一文件插入两条数据库记录。"
    "修复：loadFromDirectory() 加 synchronized，保证串行执行。"
)

# ─── 四、项目架构 ───
add_h2("四、整体架构一览")
add_code_block([
    "┌─────────────────────────────────────────────────┐",
    "│                   Vue 3 前端                     │",
    "│  ChatArea │ SettingsModal │ Sidebar │ App.vue    │",
    "│  SSE 流式渲染 │ 多模型切换 │ 主题/参数配置        │",
    "└──────────────────────┬──────────────────────────┘",
    "                       │ HTTP / SSE",
    "┌──────────────────────▼──────────────────────────┐",
    "│              Spring Boot 3.5 后端                │",
    "│  ChatService (SSE 流) │ ConversationService      │",
    "│  KnowledgeService (RAG) │ PromptTemplateService  │",
    "│  Function Calling (Tool / WebSearch / MCP)       │",
    "└────┬──────────┬──────────┬────────────┬─────────┘",
    "     │          │          │            │",
    "   MySQL     Redis      DeepSeek    Elasticsearch",
    "  (对话存储) (多级缓存)  (LLM API)   (向量+全文索引)",
])

# ─── 五、高频面试问题 ───
add_h2("五、高频面试追问 & 参考回答")

add_h3("技术深度类")

add_qa(
    "Elasticsearch 和 Milvus / Qdrant 等专用向量数据库有什么区别？",
    [
        "ES 是通用搜索引擎，支持 dense_vector 是附加能力；",
        "专用向量库（Milvus、Qdrant、Weaviate）在 HNSW、IVF 等向量索引算法上更优化，",
        "支持十亿级规模，查询 QPS 更高。",
        "本项目知识库文档量小，用 ES 同时搞定全文和向量检索，减少基础设施依赖。",
        "规模增长后可无缝迁移到专用向量库，KnowledgeService 只需替换底层实现。",
    ]
)

add_qa(
    "Spring Boot 3.x 相比 2.x 有哪些重要变化？",
    [
        "① Java 17+ 为最低要求，利用 Record、Sealed Class、Text Block 等新语法；",
        "② 底层基于 Spring 6，命名空间全面切换 javax.* → jakarta.*；",
        "③ GraalVM 原生镜像编译支持更成熟，启动时间可降至毫秒级；",
        "④ 虚拟线程（Project Loom）集成，可用 spring.threads.virtual.enabled=true 开启；",
        "⑤ 观测性（Micrometer Tracing）内置，无需额外 Sleuth。",
    ]
)

add_qa(
    "MyBatis-Plus 和 JPA / Hibernate 怎么选？",
    [
        "MyBatis-Plus：SQL 可见、灵活控制、学习成本低，适合复杂查询和对 SQL 性能敏感的场景；",
        "JPA/Hibernate：自动生成 SQL、对象关系映射更彻底，适合 CRUD 为主的业务。",
        "本项目选 MyBatis-Plus，因为对话历史查询需要分页、条件过滤，SQL 可见更便于调优。",
    ]
)

add_h3("项目经历类")

add_qa(
    "这个项目做了多久，遇到的最大挑战是什么？",
    [
        "技术挑战主要有三个：",
        "① Spring AOP 自调用失效：排查 @Transactional 不回滚的问题时，",
        "   发现根本原因是 this.method() 绕过代理，最终用 @Lazy 自注入解决；",
        "② Redis 序列化异常：List.of() 的 ImmutableCollections 序列化后反序列化崩溃，",
        "   通过 unless 条件彻底规避；",
        "③ RAG 时效性：设计 MD5 增量更新 + 时效衰减评分，在不重建全量索引的前提下",
        "   保持知识库内容的新鲜度。",
    ]
)

add_qa(
    "如果要把这个项目上生产，你会优先补哪些东西？",
    [
        "按优先级排序：",
        "① 用户认证与鉴权（JWT / OAuth2），当前无用户体系；",
        "② API 限流（Bucket4j / Redis 令牌桶），防止 LLM 费用失控；",
        "③ 监控告警（Spring Actuator + Prometheus + Grafana），对话延迟、错误率可视化；",
        "④ Docker Compose / K8s 部署配置，环境一致性保障；",
        "⑤ 对话数据加密存储，避免敏感内容明文落库。",
    ]
)

# ─── 六、一句话总结 ───
add_h2("六、简历 / 自我介绍中的一句话亮点")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(16)
# 左边框装饰
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
left = OxmlElement("w:left")
left.set(qn("w:val"), "single")
left.set(qn("w:sz"), "12")
left.set(qn("w:space"), "8")
left.set(qn("w:color"), "1A56DB")
pBdr.append(left)
pPr.append(pBdr)
run = p.add_run(
    "基于 Spring Boot 3.5 + Vue 3 的全栈 AI 对话平台，核心亮点：\n"
    "流式 SSE 响应（含中断恢复）、RAG 知识库（BM25 + kNN 混合检索 + 时效衰减）、\n"
    "Function Calling 联网搜索、Redis 分层缓存（含故障降级）、多模型动态切换。\n"
    "在开发过程中解决了 Spring AOP 自调用失效、Jackson 不可变集合序列化等工程细节问题。"
)
set_font(run, size=11.5, bold=False, italic=True)

# ── 保存 ──
output_path = r"D:\Code Practice\MySpringAIAgent\项目亮点与面试要点.docx"
doc.save(output_path)
print(f"saved: {output_path}")
